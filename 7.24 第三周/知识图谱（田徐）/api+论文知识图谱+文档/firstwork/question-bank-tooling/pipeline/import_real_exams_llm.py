"""
真题导入器 v3 —— 用大模型来读真题（这是解决"答案全有问题"的正确办法）

【为什么放弃纯规则解析】
我用正则解析这批 .doc 试了很多轮，一直有问题：
  · .doc 转换后，题干可能跨好几行，答案边界没有固定标记；
  · 答案里的"1) 2) 3)"小标号会被误认成新题号；
  · 图的题注（"图9 铣削工件表面定位示意图"）会被当成答案；
  · 公式是嵌入对象，转文本时被吞掉，留下"计算工序尺寸和。"这种断头句。
规则每修好一个，就冒出另一个。**这就是上一版答案全是垃圾的根本原因。**

【正确办法：让模型来读】
把整份卷子（带图片标记、带表格）交给大模型，让它按结构吐出 JSON。
模型天生擅长处理这种"人看得懂、正则看不懂"的排版。
而且我们让它**自己判断哪些题是残缺的**（公式丢了、图没了），残缺的直接标记丢弃，
绝不硬编。

【质量保障（三道闸）】
  1. 模型必须逐题标 `complete: true/false`（题干或答案残缺就标 false）
  2. 代码再校验一遍：答案太短、答案是图注、有断头句 -> 丢弃
  3. 需要看图的题，必须真的抽到了图，否则丢弃

用法：
    # 先干跑一份，看看质量（不写库，免费看效果）
    python3 pipeline/import_real_exams_llm.py --provider deepseek --limit 1 --dry-run

    # 正式导入（会先清掉旧的 R_ 真题）
    python3 pipeline/import_real_exams_llm.py --provider deepseek

    # 只看现在导入了多少
    python3 pipeline/import_real_exams_llm.py --stat
"""
import argparse
import glob
import hashlib
import json
import os
import re
import sqlite3
import sys
import zipfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
import sys as _qb_sys
_qb_sys.path.insert(0, BASE_DIR)
from qb_runtime import connect_database, database_path
DB_PATH = database_path()
DOCX_DIR = os.path.join(BASE_DIR, "data", "real_exams", "docx")
IMG_DIR = os.path.join(BASE_DIR, "assets", "images", "real_exams")

SCORE_RE = re.compile(r"[（(]\s*(\d+)\s*分\s*[)）]")

SYSTEM = """你是一位助教，正在把一份【机械制造工艺学】的期末试卷（含参考答案）整理成结构化题库。

【卷子的真实结构】本课程真题固定是这三种题型：
  一、名词解释（每题4分）
  二、简述题（每题10分）
  三、分析计算题（每题10~20分）

【你要做的】把卷子里的每一道题抽出来，输出 JSON。注意：

1. **题干要完整**：题干可能跨好几行，把它们拼成一句完整的题目。
   题干里的"（10分）"这种分值标注要去掉（单独放进 total_score）。

2. **答案要对得上题**：答案是紧跟在题干后面的那段内容。
   ⚠️ 千万不要把下面这些当成答案：
     - 图的题注（如"图9 铣削工件表面定位、夹紧示意图"）—— 这是图的标题，不是答案！
     - 表格的表头（如"正态分布Z/F(z)值表"）
     - 下一道题的题干

3. **采分点**：参考答案里常带"（1分）""（2分）"这样的标注，那是采分点。
   把它们抽成 rubric 列表。没有标注就留空数组。

4. **【最重要】残缺的题必须标出来，不要硬编！**
   这份文档是从 .doc 转来的，**里面的数学公式和零件图在转换时可能丢失了**。
   如果你看到下面这些迹象，说明这道题已经残缺，**必须把 complete 设为 false**：
     - 句子断头，像"计算工序尺寸和。""要求为mm。""键槽深为mm"（公式被吞了）
     - 题干说"如下图所示""分析下图"，但正文里没有图的内容
     - 答案只剩一个表头或一句图注
     - 答案是空的
   宁可标 false 丢掉，也不要编一个答案填进去。**编造答案是最严重的错误。**

【只输出一个JSON对象，不要用Markdown代码块包裹】
{
  "questions": [
    {
      "type": "名词解释 或 简述 或 计算",
      "stem": "完整题干",
      "answer": "完整参考答案",
      "total_score": 本题满分(整数),
      "rubric": [{"point": "采分点描述", "score": 分值}],
      "needs_figure": true/false,   // 这道题是否必须看图才能做
      "complete": true/false,        // 题干和答案是否都完整（残缺就false）
      "incomplete_reason": "若complete=false，说明缺什么"
    }
  ]
}"""

USER_TMPL = """下面是一份试卷的全文（从 .doc 转换而来）。
文档里检测到 {n_img} 张插图，[图片] 标记表示该位置原本有一张图。

---
{text}
---

请把这份卷子里的题目抽成 JSON。记住：残缺的题标 complete=false，不要编造答案。"""


def read_docx_marked(path):
    """读 docx，把有图的段落标上 [图片]，表格也读进来。"""
    from docx import Document
    d = Document(path)
    lines = []
    n_img = 0
    for p in d.paragraphs:
        xml = p._p.xml
        has_img = ("graphicData" in xml) or ("v:imagedata" in xml) or ("w:object" in xml)
        t = p.text.strip()
        if has_img:
            n_img += 1
            lines.append(("[图片] " + t) if t else "[图片]")
        elif t:
            lines.append(t)
    for tb in d.tables:
        rows = [" | ".join(c.text.strip() for c in r.cells)
                for r in tb.rows if any(c.text.strip() for c in r.cells)]
        if rows:
            lines.append("[表格]\n" + "\n".join(rows))
    return "\n".join(lines), n_img


def extract_media(path):
    got = []
    try:
        z = zipfile.ZipFile(path)
    except Exception:
        return got
    os.makedirs(IMG_DIR, exist_ok=True)
    for n in z.namelist():
        if not n.startswith("word/media/"):
            continue
        ext = os.path.splitext(n)[1].lower()
        if ext not in (".png", ".jpg", ".jpeg", ".gif"):
            continue
        data = z.read(n)
        if len(data) < 6000:
            continue
        fn = hashlib.md5(data).hexdigest()[:10] + ext
        fp = os.path.join(IMG_DIR, fn)
        if not os.path.exists(fp):
            with open(fp, "wb") as f:
                f.write(data)
        got.append(os.path.relpath(fp, BASE_DIR).replace("\\", "/"))
    return got


# 代码侧的第二道闸（不信模型一家之言）
TRUNC = [re.compile(r"(尺寸|直径|深度|余量|偏差|角度|长度|键槽)\s*和?\s*[。.]$"),
         re.compile(r"为\s*mm"), re.compile(r"至\s*[。.]$"), re.compile(r"[＝=]\s*[。.]"),
         re.compile(r"深为\s*mm"), re.compile(r"要求为\s*[。.]")]
FIG_CAP = re.compile(r"^\s*(题?\s*\d+\s*图|图\s*[\d\-–\.]+)")


def code_check(q, has_images):
    stem, ans = (q.get("stem") or "").strip(), (q.get("answer") or "").strip()
    if not q.get("complete", True):
        return False, "模型判定残缺：" + (q.get("incomplete_reason") or "未说明")
    if len(stem) < 4:
        return False, "题干过短"
    if len(ans) < 12:
        return False, "答案过短"
    if FIG_CAP.match(ans.split("\n")[0]) and len(ans) < 45:
        return False, "答案是图的题注"
    for p in TRUNC:
        if p.search(stem) or p.search(ans):
            return False, "有公式被吞掉的断头句"
    if q.get("needs_figure") and not has_images:
        return False, "需要看图但没有图"
    return True, ""


CHAPTER_KEYWORDS = [
    ("第二章_机械加工工艺规程设计", ["工序", "工步", "工位", "生产纲领", "工艺规程", "基准",
                                     "加工余量", "尺寸链", "加工阶段", "毛坯", "工艺路线"]),
    ("第三章_机床夹具设计", ["夹具", "定位", "自由度", "V形块", "过定位", "欠定位", "夹紧", "支承"]),
    ("第四章_机械加工精度及其控制", ["加工精度", "加工误差", "误差复映", "刚度", "回转误差",
                                     "主轴", "热变形", "正态分布", "合格品率", "工艺能力", "自激振动", "振动"]),
    ("第五章_机械加工表面质量及其控制", ["表面质量", "表面粗糙度", "冷作硬化", "残余应力",
                                          "表面层", "耐磨性", "疲劳强度", "磨削烧伤"]),
    ("第六章_机器装配工艺过程设计", ["装配", "总装", "选择装配", "分组选配", "修配", "互换"]),
    ("第七章_机械制造工艺理论和技术的发展", ["先进制造", "成组技术", "并行工程", "智能制造", "特种加工"]),
    ("第一章_绪论", ["制造技术", "制造系统", "切削", "刀具", "前角", "后角", "机床", "成型方法"]),
]


def guess_chapter(text):
    best, sc = "第一章_绪论", 0
    for ch, kws in CHAPTER_KEYWORDS:
        s = sum(1 for k in kws if k in text)
        if s > sc:
            best, sc = ch, s
    return best


def pick_files():
    """只挑标准试卷格式的（含名词解释+简述题）。
    【修bug】之前这里 except 直接 continue，如果 python-docx 没装，
    就会把 24 份文件全部静默跳过，最后打印"解析 0 份真题"——你就是这么被坑的。
    现在会把错误明确报出来。"""
    files = sorted(glob.glob(os.path.join(DOCX_DIR, "*.docx")))
    if not files:
        raise SystemExit(
            f"❌ 没找到任何真题文件。请确认这个目录里有 .docx：\n   {DOCX_DIR}")
    out, errs = [], []
    for f in files:
        try:
            txt, _ = read_docx_marked(f)
        except Exception as e:
            errs.append(f"{os.path.basename(f)}: {e}")
            continue
        head = txt[:3000]
        if re.search(r"解释下列.{0,10}名词|名词解释", head) and "简述题" in head:
            out.append(f)
    if errs:
        print(f"⚠ 有 {len(errs)} 份文件读不了，前3个原因：")
        for e in errs[:3]:
            print("   ", e)
        if len(errs) == len(files):
            raise SystemExit(
                "❌ 所有文件都读不了。最常见原因：没装 python-docx。\n"
                "   请先执行：pip install python-docx")
    print(f"共 {len(files)} 份文件，其中 {len(out)} 份是标准试卷格式")
    return out


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--provider", default=None, help="用哪个模型来读真题，如 deepseek")
    ap.add_argument("--limit", type=int, default=None, help="只处理前N份（试跑用）")
    ap.add_argument("--dry-run", action="store_true", help="不写库，只看解析质量")
    ap.add_argument("--stat", action="store_true")
    a = ap.parse_args()

    if a.stat:
        conn = connect_database()
        for r in conn.execute("SELECT question_type, COUNT(*), SUM(image_path IS NOT NULL) "
                              "FROM questions WHERE source='真题' GROUP BY question_type"):
            print(f"  {r[0]:<8} {r[1]:>3} 道（带图 {r[2] or 0} 道）")
        conn.close()
        return

    if not a.provider:
        ap.error("需要 --provider（如 deepseek）。用模型读真题才靠谱，纯规则解析已被证明会出垃圾。")

    from llm.client import get_client
    from pipeline.generate_questions import parse_llm_json
    client = get_client("concept", provider_name=a.provider)

    files = pick_files()
    if a.limit:
        files = files[:a.limit]
    print(f"用 {a.provider} 解析 {len(files)} 份真题…\n")

    kept, dropped = [], []
    for i, f in enumerate(files, 1):
        try:
            text, n_img = read_docx_marked(f)
            imgs = extract_media(f)
            data = parse_llm_json(client.chat(
                SYSTEM, USER_TMPL.format(n_img=n_img, text=text[:12000]), temperature=0.1))
            qs = data.get("questions") or []
        except Exception as e:
            print(f"[{i}/{len(files)}] 读取失败：{e}")
            continue
        ok_n = 0
        for q in qs:
            good, why = code_check(q, bool(imgs))
            if good:
                q["images"] = imgs if q.get("needs_figure") else []
                q["chapter"] = guess_chapter((q.get("stem") or "") + (q.get("answer") or "")[:200])
                kept.append(q)
                ok_n += 1
            else:
                dropped.append({"stem": (q.get("stem") or "")[:45], "reason": why})
        print(f"[{i}/{len(files)}] 抽出 {len(qs)} 道，合格 {ok_n} 道")

    # 去重
    seen, uniq = set(), []
    for q in kept:
        k = hashlib.md5((q["stem"][:40]).encode()).hexdigest()
        if k in seen:
            continue
        seen.add(k)
        uniq.append(q)

    from collections import Counter
    print(f"\n✅ 合格 {len(uniq)} 道（去重后）　❌ 丢弃 {len(dropped)} 道")
    print("题型：", dict(Counter(q["type"] for q in uniq)))
    print("章节：", dict(Counter(q["chapter"] for q in uniq)))
    print("带采分点：", sum(1 for q in uniq if q.get("rubric")))
    print("\n丢弃原因：")
    for r, n in Counter(d["reason"][:30] for d in dropped).most_common(6):
        print(f"  {n:>3}  {r}")
    print("\n样例：")
    for q in uniq[:3]:
        print(f"\n  [{q['type']}·{q.get('total_score')}分] {q['stem'][:65]}")
        print(f"    答案：{(q.get('answer') or '')[:75]}")
        if q.get("rubric"):
            print("    采分点：" + " | ".join(
                f"{p.get('point','')[:14]}({p.get('score')}分)" for p in q["rubric"][:3]))

    if a.dry_run:
        print("\n（dry-run，没写库。确认质量OK后去掉 --dry-run 正式导入。）")
        return

    conn = connect_database()
    cols = [r[1] for r in conn.execute("PRAGMA table_info(questions)")]
    for c, ddl in [("rubric_json", "TEXT"), ("total_score", "INTEGER"),
                   ("source", "TEXT DEFAULT 'AI生成'"), ("image_path", "TEXT"),
                   ("image_reviewed", "INTEGER DEFAULT 0")]:
        if c not in cols:
            conn.execute(f"ALTER TABLE questions ADD COLUMN {c} {ddl}")
    old = conn.execute("SELECT COUNT(*) FROM questions WHERE question_id LIKE 'R_%'").fetchone()[0]
    conn.execute("DELETE FROM question_knowledge_map WHERE question_id LIKE 'R_%'")
    conn.execute("DELETE FROM questions WHERE question_id LIKE 'R_%'")
    print(f"\n已清掉旧的 {old} 道真题（那批答案有问题）")

    chap_node = {}
    for (ch,) in conn.execute("SELECT DISTINCT chapter FROM knowledge_points"):
        r = conn.execute("SELECT knowledge_id FROM knowledge_points WHERE chapter=? "
                         "ORDER BY learning_order LIMIT 1", (ch,)).fetchone()
        if r:
            chap_node[ch] = r[0]

    ins = 0
    for q in uniq:
        node = chap_node.get(q["chapter"])
        if not node:
            continue
        qid = "R_" + hashlib.md5(q["stem"][:60].encode()).hexdigest()[:10]
        if conn.execute("SELECT 1 FROM questions WHERE question_id=?", (qid,)).fetchone():
            continue
        conn.execute(
            """INSERT INTO questions
               (question_id, course_chapter, source_node_id, question_type, stem,
                options_json, answer, explanation, bloom_level, generation_model,
                prompt_template_id, review_status, rubric_json, total_score, source,
                image_path, image_reviewed)
               VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (qid, q["chapter"], node, q["type"], q["stem"], None, q.get("answer"), None,
             None, "真题(教师提供)", "real_exam_llm", "已通过",
             json.dumps(q.get("rubric") or [], ensure_ascii=False),
             q.get("total_score"), "真题",
             q["images"][0] if q.get("images") else None, 1 if q.get("images") else 0))
        conn.execute("INSERT OR IGNORE INTO question_knowledge_map (question_id, knowledge_id) "
                     "VALUES (?,?)", (qid, node))
        ins += 1
    conn.commit()
    conn.close()
    print(f"已写入 {ins} 道真题")


if __name__ == "__main__":
    main()
